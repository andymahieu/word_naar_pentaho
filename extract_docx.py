# extract_docx.py
from docx import Document

# Update this if your file has a different name
input_file = "Stortformulier.docx"
output_file = "Stortformulier_extracted.txt"

try:
    doc = Document(input_file)
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("=== PARAGRAPHS ===\n")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # skip empty
                f.write(f"Line {i}: {repr(para.text)}\n")

        f.write("\n=== TABLES ===\n")
        for t_idx, table in enumerate(doc.tables):
            f.write(f"\n[Table {t_idx}]\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):  # skip empty rows
                    f.write(" | ".join(cells) + "\n")

    print(f"✅ Extraction complete! Output saved to: {output_file}")
except FileNotFoundError:
    print(f"❌ Error: '{input_file}' not found.")
except Exception as e:
    print(f"❌ Error: {e}")